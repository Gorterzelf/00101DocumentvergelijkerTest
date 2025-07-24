"""
ActiZ Document Analyzer - Document Processing Service
Enhanced version with better error handling and validation
"""

import logging
import os
from typing import Optional

import PyPDF2
from docx import Document

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Service voor het verwerken van verschillende document types
    Enhanced met betere error handling en validatie
    """

    def __init__(self):
        """Initialize with configuration"""
        self.max_file_size = int(
            os.getenv("MAX_CONTENT_LENGTH", 16 * 1024 * 1024)
        )  # 16MB
        self.supported_extensions = {"pdf", "docx", "txt"}

    def extract_text(self, file_stream, filename: str) -> Optional[str]:
        """
        Extract text from uploaded file based on extension
        Enhanced met betere validatie en error handling
        """
        try:
            # Validate filename
            if not filename or "." not in filename:
                raise ValueError("Invalid filename - geen extensie gevonden")

            file_extension = filename.lower().split(".")[-1]

            # Validate supported file type
            if file_extension not in self.supported_extensions:
                raise ValueError(
                    f"Bestandstype '{file_extension}' niet ondersteund. Gebruik: {', '.join(self.supported_extensions)}"
                )

            # Check file size (if we can)
            try:
                file_stream.seek(0, 2)  # Seek to end
                file_size = file_stream.tell()
                file_stream.seek(0)  # Reset to beginning

                if file_size > self.max_file_size:
                    raise ValueError(
                        f"Bestand te groot ({file_size // (1024*1024)}MB). Maximum: {self.max_file_size // (1024*1024)}MB"
                    )

                logger.info(
                    f"Processing {filename} ({file_size} bytes, {file_extension.upper()})"
                )
            except (OSError, AttributeError):
                # Some streams don't support seek/tell - continue anyway
                logger.info(
                    f"Processing {filename} ({file_extension.upper()}) - size unknown"
                )

            # Extract based on file type
            if file_extension == "pdf":
                return self._extract_from_pdf(file_stream)
            elif file_extension == "docx":
                return self._extract_from_docx(file_stream)
            elif file_extension == "txt":
                return self._extract_from_txt(file_stream)
            else:
                raise ValueError(
                    f"Bestandstype '{file_extension}' wordt niet ondersteund"
                )

        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            raise

    def _extract_from_pdf(self, file_stream) -> str:
        """Extract text from PDF file - Enhanced"""
        try:
            file_stream.seek(0)
            pdf_reader = PyPDF2.PdfReader(file_stream)

            # Check if PDF is encrypted
            if pdf_reader.is_encrypted:
                raise ValueError(
                    "PDF is beveiligd met een wachtwoord - kan tekst niet extraheren"
                )

            # Check number of pages
            num_pages = len(pdf_reader.pages)
            if num_pages == 0:
                raise ValueError("PDF bevat geen pagina's")

            logger.info(f"PDF heeft {num_pages} pagina(s)")

            text = ""
            pages_with_text = 0

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += page_text + "\n"
                        pages_with_text += 1
                except Exception as e:
                    logger.warning(
                        f"Kon tekst niet extraheren van pagina {page_num + 1}: {str(e)}"
                    )
                    continue

            if not text.strip():
                raise ValueError(
                    "PDF bevat geen leesbare tekst - mogelijk gescande documenten of afbeeldingen"
                )

            logger.info(
                f"Tekst geëxtraheerd van {pages_with_text}/{num_pages} pagina's ({len(text)} karakters)"
            )
            return text.strip()

        except PyPDF2.errors.PdfReadError as e:
            logger.error(f"PDF read error: {str(e)}")
            raise ValueError(f"Ongeldig PDF bestand: {str(e)}")
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise ValueError(f"Kon PDF niet lezen: {str(e)}")

    def _extract_from_docx(self, file_stream) -> str:
        """Extract text from Word document - Enhanced"""
        try:
            file_stream.seek(0)
            doc = Document(file_stream)
            text = ""
            paragraph_count = 0

            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text.strip()
                if paragraph_text:
                    text += paragraph_text + "\n"
                    paragraph_count += 1

            # Extract from tables if present
            table_count = 0
            for table in doc.tables:
                table_count += 1
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text += " | ".join(row_text) + "\n"

            if not text.strip():
                raise ValueError("Word document bevat geen tekst")

            logger.info(
                f"Word document verwerkt: {paragraph_count} paragrafen, {table_count} tabellen ({len(text)} karakters)"
            )
            return text.strip()

        except Exception as e:
            if "not a zip file" in str(e).lower() or "bad zip file" in str(e).lower():
                raise ValueError(
                    "Ongeldig Word document (.docx) - bestand is beschadigd"
                )
            logger.error(f"Word extraction error: {str(e)}")
            raise ValueError(f"Kon Word document niet lezen: {str(e)}")

    def _extract_from_txt(self, file_stream) -> str:
        """Extract text from text file - Enhanced"""
        try:
            file_stream.seek(0)
            content = file_stream.read()

            if len(content) == 0:
                raise ValueError("Tekstbestand is leeg")

            # Try different encodings
            encodings_to_try = ["utf-8", "utf-8-sig", "latin-1", "cp1252", "iso-8859-1"]
            text = None
            encoding_used = None

            for encoding in encodings_to_try:
                try:
                    text = content.decode(encoding)
                    encoding_used = encoding
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                raise ValueError(
                    "Kon tekstbestand niet decoderen - onbekende karakterset"
                )

            if not text.strip():
                raise ValueError("Tekstbestand bevat geen inhoud")

            logger.info(
                f"Tekstbestand gelezen met {encoding_used} encoding ({len(text)} karakters)"
            )
            return text.strip()

        except Exception as e:
            logger.error(f"Text extraction error: {str(e)}")
            raise ValueError(f"Kon tekstbestand niet lezen: {str(e)}")

    def validate_file(self, filename: str) -> dict:
        """
        Validate file without processing - useful for frontend validation
        """
        try:
            if not filename or "." not in filename:
                return {
                    "valid": False,
                    "error": "Ongeldige bestandsnaam - geen extensie gevonden",
                }

            file_extension = filename.lower().split(".")[-1]

            if file_extension not in self.supported_extensions:
                return {
                    "valid": False,
                    "error": f"Bestandstype '{file_extension}' niet ondersteund",
                    "supported_types": list(self.supported_extensions),
                }

            return {
                "valid": True,
                "file_type": file_extension,
                "message": f"{file_extension.upper()} bestand is ondersteund",
            }

        except Exception as e:
            return {"valid": False, "error": f"Validatie fout: {str(e)}"}
