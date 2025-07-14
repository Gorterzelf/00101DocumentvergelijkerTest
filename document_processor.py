"""
Document Processor - Clean implementation with Strategy Pattern
Scheidt text extractie van analyse logica
"""

import os
import logging
from typing import Dict, Optional
from datetime import datetime
from docx import Document
from openai import AzureOpenAI
import tiktoken

from analysis_strategies import AnalysisFactory

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles document processing with clean separation of concerns"""
    
    def __init__(self):
        """Initialize the DocumentProcessor with Azure OpenAI client"""
        self.api_key = os.getenv("AZURE_OPENAI_API_KEY")
        self.endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
        self.deployment_name = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME", "gpt-4")
        self.api_version = os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-15-preview")
        
        if not all([self.api_key, self.endpoint]):
            raise ValueError("Azure OpenAI credentials not properly configured")
        
        self.client = AzureOpenAI(
            api_key=self.api_key,
            api_version=self.api_version,
            azure_endpoint=self.endpoint
        )
        
        logger.info(f"Azure OpenAI client initialized - Endpoint: {self.endpoint}")
        
        # Initialize tokenizer
        self.encoding = tiktoken.encoding_for_model("gpt-4")
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text from various file types"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        extractors = {
            '.docx': self._extract_from_docx,
            '.doc': self._extract_from_docx,
            '.pdf': self._extract_from_pdf,
            '.txt': self._extract_from_txt
        }
        
        extractor = extractors.get(file_extension)
        if not extractor:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        try:
            text = extractor(file_path)
            logger.info(f"Extracted {len(text)} characters from {os.path.basename(file_path)}")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from Word document"""
        doc = Document(file_path)
        full_text = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(' | '.join(row_text))
        
        return '\n\n'.join(full_text)
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF with fallback options"""
        # Try PyPDF2 first
        try:
            import PyPDF2
            text = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                logger.info(f"PDF has {num_pages} pages")
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text.strip():
                        text.append(page_text)
                
                if text:
                    return '\n\n'.join(text)
                else:
                    logger.warning("PyPDF2 extracted no text, trying pdfplumber")
                    
        except Exception as e:
            logger.warning(f"PyPDF2 failed: {e}, trying pdfplumber")
        
        # Fallback to pdfplumber
        try:
            import pdfplumber
            text = []
            
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
                        logger.debug(f"Page {i+1}: {len(page_text)} chars")
            
            if text:
                return '\n\n'.join(text)
            else:
                raise ValueError("No text could be extracted from PDF")
                
        except ImportError:
            raise ImportError("Install PyPDF2 or pdfplumber: pip install PyPDF2 pdfplumber")
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from plain text file with encoding detection"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        
        raise ValueError(f"Could not decode text file with any common encoding")
    
    def process_documents(self, file1_path: str, file2_path: str, 
                         file3_path: Optional[str] = None, mode: str = "version_comparison") -> Dict:
        """Process documents using the appropriate analysis strategy"""
        try:
            # Extract text from files
            logger.info(f"Processing documents in {mode} mode")
            text1 = self.extract_text_from_file(file1_path)
            text2 = self.extract_text_from_file(file2_path)
            text3 = self.extract_text_from_file(file3_path) if file3_path else None
            
            # Create appropriate strategy
            strategy = AnalysisFactory.create_strategy(mode, self.client)
            
            # Run analysis
            analysis_result = strategy.analyze(text1, text2, text3)
            
            # Add metadata
            result = {
                'mode': mode,
                'timestamp': datetime.now().isoformat(),
                'doc1_name': os.path.basename(file1_path),
                'doc2_name': os.path.basename(file2_path),
                'doc3_name': os.path.basename(file3_path) if file3_path else None,
                'doc1_length': len(text1),
                'doc2_length': len(text2),
                'doc3_length': len(text3) if text3 else 0,
                **analysis_result
            }
            
            logger.info(f"Analysis completed successfully for {mode}")
            return result
            
        except Exception as e:
            logger.error(f"Error processing documents: {e}")
            raise
    
    def format_results_for_display(self, results: Dict) -> Dict:
        """Format results for web display using strategy pattern"""
        try:
            mode = results.get('mode')
            strategy = AnalysisFactory.create_strategy(mode, self.client)
            
            # Get strategy-specific formatting
            formatted = strategy.format_results(results)
            
            # Add common metadata
            formatted['metadata'] = {
                'mode': mode,
                'timestamp': results.get('timestamp', ''),
                'doc1_name': results.get('doc1_name', ''),
                'doc2_name': results.get('doc2_name', ''),
                'doc3_name': results.get('doc3_name', ''),
                'doc1_length': f"{results.get('doc1_length', 0):,} karakters",
                'doc2_length': f"{results.get('doc2_length', 0):,} karakters",
                'doc3_length': f"{results.get('doc3_length', 0):,} karakters" if results.get('doc3_length') else None
            }
            
            return formatted
            
        except Exception as e:
            logger.error(f"Error formatting results: {e}")
            return {
                'metadata': {'mode': 'error'},
                'error': str(e)
            }